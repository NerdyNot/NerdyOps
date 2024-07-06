import pypandoc
from docx import Document
from PyPDF2 import PdfReader
import os


def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = []
    for page in reader.pages:
        text.append(page.extract_text())
    return "\n".join(text)

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return "\n".join(text)

def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text_from_md(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        md_content = file.read()
    return md_content

def convert_pdf_to_word(pdf_path):
    word_path = f"{os.path.splitext(pdf_path)[0]}.docx"
    pypandoc.convert_file(pdf_path, 'docx', outputfile=word_path)
    return word_path

def convert_word_to_pdf(word_path, pdf_path):
    pypandoc.convert_file(word_path, 'pdf', outputfile=pdf_path)

def create_translated_word(original_word_path, translated_paragraphs, output_word_path):
    doc = Document(original_word_path)
    for i, para in enumerate(doc.paragraphs):
        if i < len(translated_paragraphs):
            para.text = translated_paragraphs[i]
    doc.save(output_word_path)